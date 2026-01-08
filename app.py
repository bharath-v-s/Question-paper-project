from flask import Flask, render_template, request, jsonify, redirect, send_file, url_for, session,flash
from flask_sqlalchemy import SQLAlchemy
import pandas as pd
import random
import os

from docx import Document
from docx.shared import Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://root:password@localhost/qp_generator'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = "university_secret_key" # Required for sessions
app.config['UPLOAD_FOLDER'] = 'uploads'



db = SQLAlchemy(app)

# --- MODELS ---
class School(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100))
    depts = db.relationship('Department', backref='school', lazy=True)

class Department(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    school_id = db.Column(db.Integer, db.ForeignKey('school.id'))
    name = db.Column(db.String(100))
    level = db.Column(db.String(5))
    subjects = db.relationship('Subject', backref='dept', lazy=True)
    
class GridType(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50)) # e.g., "Theory Only" or "Theory & Problem"
    has_problem_column = db.Column(db.Boolean, default=False)

# Update Subject to link to a GridType
class Subject(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    dept_id = db.Column(db.Integer, db.ForeignKey('department.id'))
    grid_type_id = db.Column(db.Integer, db.ForeignKey('grid_type.id')) # New Link
    name = db.Column(db.String(100))
    code = db.Column(db.String(20))
    semester = db.Column(db.Integer)
    pattern_name = db.Column(db.String(50), default="Pattern_1")
    

class ExamPattern(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), unique=True, nullable=False) # e.g., "Pattern_1"
    total_marks = db.Column(db.Integer, default=100)
    sections = db.relationship('PatternSection', backref='pattern', lazy=True, cascade="all, delete-orphan")

class PatternSection(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    pattern_id = db.Column(db.Integer, db.ForeignKey('exam_pattern.id'), nullable=False)
    section_name = db.Column(db.String(10), nullable=False) # e.g., "SecA"
    count = db.Column(db.Integer, nullable=False)
    total_in_paper = db.Column(db.Integer, nullable=False) # The 'total' field in your dict
    marks = db.Column(db.Integer, nullable=False)
    note = db.Column(db.Text)



if not os.path.exists('uploads'):
    os.makedirs('uploads')

question_bank_df = None
def load_question_bank(filepath):
    """Loads the cleaned question bank without skipping rows."""
    global question_bank_df
    try:
        if filepath.endswith('.csv'):
            df = pd.read_csv(filepath)
        elif filepath.endswith('.xlsx'):
            df = pd.read_excel(filepath)
        
        df.columns = df.columns.str.strip()
        # ['Unit', 'Marks','K Level', 'Question'] is mandatory in the uploaded file
        required = ['Unit', 'Marks','K Level', 'Question'] 
        if not all(col in df.columns for col in required):
            return False
            
        # Standardize the 'Type' column
        if 'Type' not in df.columns:
            # Create the column if it's missing
            df['Type'] = 'Theory'
        else:
            # Clean existing values
            df['Type'] = df['Type'].fillna('Theory').astype(str).str.strip()
            
        # Clean other columns
        df['Question'] = df['Question'].astype(str).str.strip()
        df['Marks'] = pd.to_numeric(df['Marks'], errors='coerce').fillna(0).astype(int)
        df['Unit'] = pd.to_numeric(df['Unit'], errors='coerce').fillna(0).astype(int)
        df['K Level'] = df['K Level'].astype(str).str.strip().str.upper()
        
        df = df.dropna(subset=['Marks', 'Unit', 'Question', 'K Level'])
        question_bank_df = df
        return True
    except Exception as e:
        print(f"Error loading file: {e}")
        return False
    
def sample_from_unit(unit_no, marks, count, q_type=None):
    """Filters bank by Unit, Marks, and optionally Theory/Problem type."""
    global question_bank_df
    if question_bank_df is None or count <= 0:
        return []

    # 1. Ensure the DataFrame columns are numeric to match the function arguments
    # 2. Strip whitespace from the Type column for clean comparison
    temp_df = question_bank_df.copy()
    temp_df['Marks'] = pd.to_numeric(temp_df['Marks'], errors='coerce')
    temp_df['Unit'] = pd.to_numeric(temp_df['Unit'], errors='coerce')
    temp_df['Type'] = temp_df['Type'].astype(str).str.strip()

    # Filter by Unit and Marks
    mask = (temp_df['Unit'] == int(unit_no)) & (temp_df['Marks'] == int(marks))
    
    # Optional Type filtering (Theory/Problem) 
    if q_type and 'Type' in temp_df.columns:
        # Make the comparison case-insensitive and remove spaces
        mask = mask & (temp_df['Type'].astype(str).str.strip().str.lower() == q_type.lower().strip())
    
    pool = temp_df[mask]
    
    # DEBUG: See why Section B/C might be empty
    print(f"Pool size for Unit {unit_no}, {marks}M: {len(pool)}") 
    
    if len(pool) == 0:
        return [] # No questions found matching these exact criteria
    
    if len(pool) < count:
        return pool.to_dict('records') # Return all available if fewer than requested
    
    return pool.sample(n=count).to_dict('records')


# --- EXAM PATTERNS DEFINED AS DICT (for initial migration) ---
# remove the triple quotes below after migration is done
"""# --- PATTERNS DATA ---
EXAM_PATTERNS = {
    "Pattern_1": {
        "SecA": {"count": 10, "total": 12, "marks": 3, "note": "Answer any ten questions"},
        "SecB": {"count": 5, "total": 7, "marks": 6, "note": "Answer any five questions"},
        "SecC": {"count": 4, "total": 6, "marks": 10, "note": "Answer any four questions"},
        "total_marks": 100
    },"Pattern_2": {
        "SecA": {"count": 10, "total": 12, "marks": 2, "note": "Answer any Ten questions"},
        "SecB": {"count": 5, "total": 8, "marks": 8, "note": "Answer any five questions"},
        "SecC": {"count": 2, "total": 3, "marks": 20, "note": "Answer any two questions"},
        "total_marks": 100
    },"Pattern_3": {
        "SecA": {"count": 10, "total": 10, "marks": 2, "note": "Answer all questions"},
        "SecB": {"count": 5, "total": 8, "marks": 8, "note": "Answer any five questions"},
        "SecC": {"count": 2, "total": 3, "marks": 20, "note": "Answer any two questions"},
        "total_marks": 100
    },"Pattern_4": {
        "SecA": {"count": 5, "total": 7, "marks": 2, "note": "Answer any five questions"},
        "SecB": {"count": 4, "total": 6, "marks": 5, "note": "Answer any four questions"},
        "SecC": {"count": 2, "total": 2, "marks": 20, "note": "Answer any two questions(Internal Choice)"},
        "total_marks": 50
    },"Pattern_5": {
        "SecA": {"count": 20, "total": 20, "marks": 1, "note": "Answer all questions"},
        "SecB": {"count": 5, "total": 8, "marks": 7, "note": "Answer any five questions"},
        "SecC": {"count": 3, "total": 5, "marks": 15, "note": "Answer any three questions"},
        "total_marks": 100
    }

}"""
# Predefined Exam Patterns (moved to DB migration section)
def get_pattern_dict(pattern_name):
    pattern = ExamPattern.query.filter_by(name=pattern_name).first()
    if not pattern:
        return None
    
    pattern_dict = {"total_marks": pattern.total_marks}
    for sec in pattern.sections:
        pattern_dict[sec.section_name] = {
            "count": sec.count,
            "total": sec.total_in_paper,
            "marks": sec.marks,
            "note": sec.note
        }
    return pattern_dict

# --- ROUTES ---
@app.route('/')
def index():
    schools = School.query.all()
    return render_template('index.html', schools=schools)

@app.route('/get_departments/<int:sid>/<level>')
def get_depts(sid, level):
    #filters departments by both school_id and level
    depts = Department.query.filter_by(school_id=sid, level=level).all()
    return jsonify([{'id': d.id, 'name': d.name} for d in depts])

@app.route('/get_subjects/<int:dept_id>/<int:semester>') # Added <int:semester> to the URL
def get_subjects(dept_id, semester):
    #filters subjects by both dept_id and semester
    subjects = Subject.query.filter_by(dept_id=dept_id, semester=semester).all()
    return jsonify([{'id': s.id, 'name': s.name} for s in subjects])

@app.route('/get_pattern_details/<int:dept_id>/<int:subject_id>')
def get_pattern_details(dept_id, subject_id):
    # Use modern SQLAlchemy Session.get()
    subject = db.session.get(Subject, subject_id)
    if not subject:
        return jsonify({"error": "Subject not found"}), 404

    # Fetch Grid Configuration
    grid_config = db.session.get(GridType, subject.grid_type_id)
    
    # Fetch pattern from Database instead of EXAM_PATTERNS dict
    pattern_rules = get_pattern_dict(subject.pattern_name or "Pattern_1")
    
    return jsonify({
        "config": pattern_rules,
        "is_split": grid_config.has_problem_column if grid_config else False,
        "pattern_id": subject.pattern_name or "Pattern_1"
    })

#uploads will create a new folder named uploads in the project directory to store uploaded files
@app.route('/upload', methods=['POST'])
def upload():
    file = request.files.get('file')
    if file and file.filename != '':
        path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(path)
        
        # Uses your helper function to load data into question_bank_df
        success = load_question_bank(path) 
        
        if success:
            # return success response in upload in index.html
            return jsonify({
                "status": "success", 
                "message": f"Question Bank '{file.filename}' loaded successfully!"
            })
        else:
            return jsonify({
                "status": "error", 
                "message": "Missing required columns: Sl. No, Unit, Marks, K Level, Question"
            }), 400
            
    return jsonify({"status": "error", "message": "No file selected"}), 400

@app.route('/generate', methods=['POST'])
def generate():
    global question_bank_df
    if question_bank_df is None:
        flash("Please upload a Question Bank first!", "danger")
        return redirect(url_for('index'))

    # 1. Get Subject details to find the correct Pattern
    subject_id = request.form.get('subject_id')
    subject_obj = db.session.get(Subject, subject_id)

    # Strictly store these in session for the download and swap routes
    session['selected_subject_id'] = subject_id
    session['selected_subject_name'] = subject_obj.name if subject_obj else "Unknown Subject"
    
    if not subject_obj:
        flash("Invalid Subject selected.", "danger")
        return redirect(url_for('index'))

    # 2. Identify the pattern (e.g., 'Pattern_1')
    pattern_key = subject_obj.pattern_name if subject_obj.pattern_name else "Pattern_1"
    pattern_data = get_pattern_dict(pattern_key)
    selected_questions = []

    # 3. Process the grid inputs
    for key, value in request.form.items():
        # Keys from HTML are: u1_SecA_t, u1_SecB_t, etc.
        if value and value.isdigit() and int(value) > 0:
            parts = key.split('_') 
            if len(parts) < 3: continue
            
            unit_no = int(parts[0][1:]) # e.g., 1 from 'u1'
            section = parts[1]           # e.g., 'SecA', 'SecB', or 'SecC'
            count = int(value)
            q_type = "Theory" if parts[2] == 't' else "Problem"

            # DYNAMIC MARK LOOKUP: Get the marks for this specific section from the pattern
            # If section is 'SecB', it looks at pattern_data['SecB']['marks']
            section_config = pattern_data.get(section)
            if not section_config:
                continue
                
            marks = section_config.get('marks')

            # 4. Sample from the global dataframe
            qs = sample_from_unit(unit_no, marks, count, q_type)
            selected_questions.extend(qs)
            print(f"Searching for: Unit {unit_no}, Marks {marks}, Type {q_type}, Count {count}")

    session['current_paper'] = selected_questions
    session.modified = True
    return render_template('review.html', 
                           questions=selected_questions, 
                           subject=subject_obj.name, 
                           qtype="Mixed Pattern")

@app.route('/swap/<int:index>', methods=['POST'])
def swap(index):
    global question_bank_df
    paper = session.get('current_paper')
    
    if question_bank_df is None or not paper:
        return jsonify({"error": "Bank or paper missing"}), 400

    old_q = paper[index]
    current_type = old_q.get('Type', 'Theory')

    # Filter for a replacement in the global dataframe
    mask = (question_bank_df['Unit'] == int(old_q['Unit'])) & \
           (question_bank_df['Marks'] == int(old_q['Marks'])) & \
           (question_bank_df['Type'].str.strip().str.lower() == current_type.lower().strip()) & \
           (~question_bank_df['Question'].isin([q['Question'] for q in paper]))

    pool = question_bank_df[mask]
    
    if not pool.empty:
        new_q = pool.sample(n=1).to_dict('records')[0]
        # Update session data
        paper[index] = new_q
        session['current_paper'] = paper
        session.modified = True
        # Return only the new question data
        return jsonify({"success": True,
                        "new_question": new_q['Question'],
                        "new_klevel": new_q['K Level'] # Return new K-Level for AJAX
                        })
    
    return jsonify({"success": False, "message": "No alternative found"}), 404

@app.route('/download/docx', methods=['POST', 'GET'])
def download_docx():
    # 1. Retrieve the exact questions from the current session
    questions = session.get('current_paper')
    subject_name = session.get('selected_subject_name')
    subject_id = session.get('selected_subject_id')

    if not questions:
        return "No questions found", 400
    
    subject_obj = db.session.get(Subject, subject_id) if subject_id else None
    pattern_key = subject_obj.pattern_name if subject_obj else "Pattern_1"
    pattern_data = get_pattern_dict(pattern_key)
    # 2. Convert to DataFrame to handle grouping and sorting
    df_paper = pd.DataFrame(questions)
    
    doc = Document()
    # --- 1. Layout: Margins ---
    # Top: 1.27cm, Bottom: 1.27cm, Right: 1.27cm, Left: 1.5cm
    section = doc.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.left_margin = Cm(1.5)
    # --- Header: REG NO Box ---
    reg_p = doc.add_paragraph()
    reg_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    reg_p.paragraph_format.space_after = Pt(0) # 0pt spacing
    reg_p.add_run("REG. NO : _______________________").bold = True

    # --- Header: College Name ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run('Weightage based QUESTION PAPER generator')
    run.font.size = Pt(13)
    run.bold = True
    
    exam_info = doc.add_paragraph()
    exam_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exam_info.paragraph_format.space_after = Pt(0)
    # --- exam info (static)---
    exam_info.add_run('NOV 2025\n').bold = True
    # --- exam info (dynamic subject name and code)---
    #exam_info.add_run(f"{subject_name if subject_name else 'SUBJECT NAME'}\
    exam_info.add_run(f"{subject_obj.code if subject_obj else 'CODE'}").bold = True

    # --- Header: Marks & Time ---
    marks_time = doc.add_paragraph()
    marks_time.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    marks_time.paragraph_format.space_after = Pt(12) # Space before Section A
    marks_time.add_run(f'MAX. MARKS: {pattern_data.get("total_marks", 100)}\nTIME: 3 HRS.\t').bold = True

    # --- Dynamic Sections ---
    global_q_num = 1
    # Sorted section keys: SecA, SecB, SecC
    section_keys = sorted([k for k in pattern_data.keys() if k.startswith('Sec')])

    for sec_key in section_keys:
        sec_cfg = pattern_data[sec_key]
        m = sec_cfg['marks']
        count = sec_cfg['count']
        total_sec_marks = count * m
        label = sec_key[-1] # Extracts 'A' from 'SecA'

        # Filter questions matching the marks for this section
        sec_qs = df_paper[df_paper['Marks'] == m]
        
        if not sec_qs.empty:
            # Modified Section Heading per your request
            sec_head = doc.add_paragraph()
            sec_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sec_head.paragraph_format.space_after = Pt(0) # 0pt spacing
            sec_head.paragraph_format.left_indent = Pt(0)
            sec_head.paragraph_format.right_indent = Pt(0)

            header_text = f"SECTION – {label} ({count} X {m} = {total_sec_marks} MARKS)"
            sec_head.add_run(header_text).bold = True
            
            # Dynamic Note: Centered and Italicized
            sec_note = doc.add_paragraph()
            sec_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sec_note.paragraph_format.space_after = Pt(6) # 6pt space before questions start

            note_text = f"({sec_cfg.get('note', 'Answer as required')})"
            sec_note.add_run(note_text).italic = True

            # Question Listing with Continuous Numbering
            for _, row in sec_qs.iterrows():
                q_p = doc.add_paragraph()
                q_p.paragraph_format.left_indent = Cm(1) # Left indent 1cm
                q_p.paragraph_format.right_indent = Cm(0)
                q_p.paragraph_format.space_after = Pt(6) # Spacing 6pt

                q_p.add_run(f"{global_q_num}. {row['Question']}")
                global_q_num += 1

    # --- Footer ---
    footer = doc.add_paragraph('\n******')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], "Final_Question_Paper.docx")
    doc.save(file_path)
    return send_file(file_path, as_attachment=True)
           
    
# --- INITIALIZATION ---
if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        if not School.query.first():
            s1 = School(name="School of IT")
            db.session.add(s1)
            db.session.commit()
            db.session.add(Department(school_id=s1.id, name="BCA", level="UG"))
            db.session.commit()
            #remove below code after first run(no need to migrate multiple times(it will duplicate data or cause errors))
        """if not ExamPattern.query.first():
            print("Migrating patterns to database...")
            for p_name, p_data in EXAM_PATTERNS.items():
                new_pattern = ExamPattern(name=p_name, total_marks=p_data['total_marks'])
                db.session.add(new_pattern)
                db.session.flush() # Get the ID for section mapping

                for sec_key, sec_val in p_data.items():
                    if sec_key.startswith('Sec'):
                        new_section = PatternSection(
                            pattern_id=new_pattern.id,
                            section_name=sec_key,
                            count=sec_val['count'],
                            total_in_paper=sec_val['total'],
                            marks=sec_val['marks'],
                            note=sec_val['note']
                        )
                        db.session.add(new_section)
            
            db.session.commit()
            print("Pattern migration complete.")"""

    app.run(debug=True)